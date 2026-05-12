#!/usr/bin/env python3
"""Generate SOP Word Document for Accountants."""

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE
from pathlib import Path
import datetime

doc = Document()

# ── Page Setup ─────────────────────────────────────────────────────
for section in doc.sections:
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

# ── Styles ─────────────────────────────────────────────────────────
style = doc.styles['Normal']
font = style.font
font.name = 'Calibri'
font.size = Pt(11)

# Custom heading colors
for i in range(1, 4):
    h = doc.styles[f'Heading {i}']
    h.font.color.rgb = RGBColor(26, 35, 126)  # Dark blue
    h.font.name = 'Calibri'


def add_table(data, col_widths=None, header=True):
    """Add a formatted table."""
    table = doc.add_table(rows=len(data), cols=len(data[0]))
    table.style = 'Light Grid Accent 1'
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    for i, row_data in enumerate(data):
        for j, cell_text in enumerate(row_data):
            cell = table.cell(i, j)
            cell.text = str(cell_text)
            for paragraph in cell.paragraphs:
                paragraph.style.font.size = Pt(10)
                for run in paragraph.runs:
                    run.font.size = Pt(10)
                    run.font.name = 'Calibri'
                    if i == 0 and header:
                        run.bold = True
    return table


def add_note(text, style="info"):
    """Add a highlighted note box."""
    p = doc.add_paragraph()
    prefix = {"info": "ℹ️ ", "warning": "⚠️ ", "important": "❗ ", "tip": "💡 "}
    run = p.add_run(prefix.get(style, "") + text)
    run.font.size = Pt(10)
    run.italic = True
    run.font.color.rgb = RGBColor(100, 100, 100)


# ════════════════════════════════════════════════════════════════════
#  COVER PAGE
# ════════════════════════════════════════════════════════════════════

for _ in range(6):
    doc.add_paragraph()

title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run("THOTA HOSPITALITY LLP")
run.bold = True
run.font.size = Pt(28)
run.font.color.rgb = RGBColor(26, 35, 126)

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = subtitle.add_run("Tally Expense Upload\nStandard Operating Procedure")
run.font.size = Pt(18)
run.font.color.rgb = RGBColor(80, 80, 80)

doc.add_paragraph()

meta = doc.add_paragraph()
meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = meta.add_run(f"Version 1.0  |  {datetime.date.today().strftime('%B %d, %Y')}\nFor Internal Use — Accounts Team")
run.font.size = Pt(11)
run.font.color.rgb = RGBColor(120, 120, 120)

doc.add_page_break()


# ════════════════════════════════════════════════════════════════════
#  TABLE OF CONTENTS
# ════════════════════════════════════════════════════════════════════

doc.add_heading("Table of Contents", level=1)
toc_items = [
    "1. Overview",
    "2. Prerequisites & One-Time Setup",
    "3. How to Fill the Excel Template",
    "    3.1 Column Descriptions",
    "    3.2 Voucher Type Rules",
    "    3.3 GST Treatment",
    "    3.4 Vendor Name Rules",
    "4. Expense Ledger List (Tally Heads)",
    "5. Vendor / Party Ledger List",
    "6. Department List",
    "7. How to Upload — Step by Step",
    "8. How to Verify in Tally",
    "9. Common Errors & Fixes",
    "10. Frequently Asked Questions",
]
for item in toc_items:
    p = doc.add_paragraph(item)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.space_before = Pt(0)

doc.add_page_break()


# ════════════════════════════════════════════════════════════════════
#  1. OVERVIEW
# ════════════════════════════════════════════════════════════════════

doc.add_heading("1. Overview", level=1)

doc.add_paragraph(
    "This tool automates the entry of daily expenses into Tally Prime for THOTA HOSPITALITY LLP. "
    "Instead of manually typing each expense entry in Tally, you fill an Excel spreadsheet and "
    "upload it through a simple web application. The tool validates your data, converts it to "
    "Tally-compatible format, and pushes all entries to Tally in seconds."
)

doc.add_heading("What it does:", level=3)
bullets = [
    "Reads your Excel expense file",
    "Validates all entries (checks for errors like missing dates, invalid ledger names)",
    "Shows you a preview of all entries before uploading",
    "Converts entries to Tally vouchers (Journal & Purchase)",
    "Uploads directly to Tally Prime in one click",
]
for b in bullets:
    doc.add_paragraph(b, style='List Bullet')

doc.add_heading("Key Numbers:", level=3)
add_table([
    ["Metric", "Value"],
    ["Monthly entries handled", "~1000 per month"],
    ["Time saved per month", "~40 hours (vs manual entry)"],
    ["Supported voucher types", "Journal, Purchase"],
    ["Company", "THOTA HOSPITALITY LLP"],
    ["GST status", "5% GST holder — no ITC claimed"],
])


doc.add_page_break()

# ════════════════════════════════════════════════════════════════════
#  2. PREREQUISITES & ONE-TIME SETUP
# ════════════════════════════════════════════════════════════════════

doc.add_heading("2. Prerequisites & One-Time Setup", level=1)

doc.add_heading("What You Need:", level=2)
add_table([
    ["Requirement", "Details"],
    ["Computer", "Windows PC with the AGENT folder"],
    ["Tally Prime", "Must be running on the same computer with THOTA HOSPITALITY LLP company open"],
    ["Tally XML Server", "Must be enabled (see Tally setup below)"],
    ["Python", "Must be installed (one-time, already done on this computer)"],
    ["Browser", "Chrome, Edge, or Firefox"],
])

doc.add_paragraph()
doc.add_heading("Enable Tally XML Server (one-time):", level=2)

tally_steps = [
    "Open Tally Prime and load THOTA HOSPITALITY LLP company",
    "Press F12 (Configuration)",
    "Go to Connectivity",
    "Set \"Tally Prime Server\" → Yes",
    "Set \"Port\" → 9000",
    "Accept and restart Tally",
]
for i, step in enumerate(tally_steps, 1):
    doc.add_paragraph(f"Step {i}: {step}", style='List Number')

add_note("This only needs to be done once. After enabling, Tally will remember this setting.")

doc.add_paragraph()
doc.add_heading("AGENT Folder Location:", level=2)
p = doc.add_paragraph()
run = p.add_run("D:\\RaviGowda\\Thota Documents\\AGENT\\")
run.bold = True
run.font.size = Pt(12)
run.font.name = 'Consolas'

doc.add_page_break()


# ════════════════════════════════════════════════════════════════════
#  3. HOW TO FILL THE EXCEL TEMPLATE
# ════════════════════════════════════════════════════════════════════

doc.add_heading("3. How to Fill the Excel Template", level=1)

doc.add_paragraph(
    "Download or generate the template from the web app (\"Download Template\" button). "
    "The template has dropdown lists for most columns, so you just select values. "
    "Fill one row per expense entry."
)

doc.add_heading("3.1 Column Descriptions", level=2)

add_table([
    ["Column", "Header", "Required?", "Description", "Example"],
    ["A", "Date", "YES", "Date of expense (DD-MM-YYYY)", "01-03-2026"],
    ["B", "Voucher_Type", "YES", "Journal (no GST bill) or Purchase (GST bill)", "Journal"],
    ["C", "Department", "No", "Which department incurred the expense", "Thota Kitchen"],
    ["D", "Description", "YES", "What was the expense for", "Vegetables from local market"],
    ["E", "Expense_Ledger", "YES", "Tally expense head (SELECT FROM DROPDOWN)", "Thota Kitchen"],
    ["F", "Vendor_Name", "YES", "Tally party name (SELECT FROM DROPDOWN)", "VENDOR for Non GST Purchases"],
    ["G", "Amount", "YES", "Total amount paid (including GST if any)", "1500"],
    ["H", "GST_Rate", "No", "GST % on the bill (0, 5, 12, 18, 28)", "5"],
    ["I", "GSTIN", "No", "Vendor's GSTIN number", "29AAFCG9846E1Z7"],
    ["J", "Bill_Number", "No", "Invoice / bill number from vendor", "INV-1234"],
    ["K", "Paid_By", "No", "Person who made the payment", "Niveditha"],
    ["L", "Payment_Mode", "No", "Cash, Bank, or UPI", "UPI"],
    ["M", "Comment", "No", "Any additional notes", "Anitha vegetable shop"],
])

doc.add_paragraph()
add_note("Columns marked YES are mandatory. The upload will fail if these are missing.", "important")

doc.add_paragraph()
doc.add_heading("3.2 Voucher Type Rules", level=2)

doc.add_paragraph("Choose the voucher type based on whether you have a GST bill or not:")

add_table([
    ["Voucher Type", "When to Use", "Vendor_Name Column"],
    ["Journal", "No GST bill / petty cash / non-GST vendor", "Use \"VENDOR for Non GST Purchases\" or \"Cash\""],
    ["Purchase", "GST invoice available from registered vendor", "Use exact vendor name from Tally (e.g., \"Meat Shop\", \"Petrol Bunk\")"],
])

doc.add_paragraph()

doc.add_heading("Accounting Impact:", level=3)
add_table([
    ["Type", "Debit (Expense)", "Credit (Vendor/Cash)"],
    ["Journal", "Expense Ledger ← full amount", "VENDOR for Non GST / Cash ← full amount"],
    ["Purchase", "Expense Ledger ← full amount (incl. GST)", "Vendor party ← full amount"],
])

doc.add_paragraph()
doc.add_heading("3.3 GST Treatment", level=2)

p = doc.add_paragraph()
run = p.add_run("IMPORTANT: ")
run.bold = True
run.font.color.rgb = RGBColor(198, 40, 40)
run = p.add_run(
    "THOTA HOSPITALITY LLP is a 5% GST holder and does NOT claim input tax credit (ITC). "
    "Therefore, the full GST-inclusive amount goes to the expense ledger. "
    "There is NO CGST/SGST split in the voucher entries."
)

doc.add_paragraph(
    "The GST_Rate and GSTIN columns are for your record-keeping only. "
    "They do not affect the Tally entry. You can leave them blank if you want."
)

doc.add_paragraph()
doc.add_heading("3.4 Vendor Name Rules", level=2)

add_table([
    ["Scenario", "What to Put in Vendor_Name"],
    ["No GST bill (vegetables, flowers, petty cash)", "VENDOR for Non GST Purchases"],
    ["Cash payment, no specific vendor", "Cash"],
    ["Bank direct debit", "Axis Bank or ICICI Bank"],
    ["GST vendor with bill (has GSTIN)", "Exact Tally ledger name — e.g., Meat Shop, Petrol Bunk, The Big Market"],
])

add_note("Vendor names must match EXACTLY with Tally. Use the dropdown in the Excel template to avoid typos.", "warning")

doc.add_page_break()


# ════════════════════════════════════════════════════════════════════
#  4. EXPENSE LEDGER LIST
# ════════════════════════════════════════════════════════════════════

doc.add_heading("4. Expense Ledger List (Tally Heads)", level=1)

doc.add_paragraph(
    "Use these exact names in the Expense_Ledger column (Column E). "
    "These are available as dropdowns in the Excel template."
)

expense_ledgers = [
    ("Thota Kitchen", "All kitchen expenses — vegetables, chicken, grocery, milk, cooking supplies"),
    ("Thota Decor", "Decoration items — flowers, gypsy, event decorations"),
    ("Studio Decor", "Studio-specific decoration and supplies"),
    ("Fuel Expenses", "Petrol, diesel for vehicles"),
    ("Transportation Expenses", "Porter, cab, auto charges"),
    ("Staff Welfare", "Staff food, staff vegetable, eggs for staff, staff welfare items"),
    ("Staffwelfare-Studio", "Studio staff welfare expenses"),
    ("Repairs & Maintenance", "Plumbing, electrical, AC repair, general maintenance"),
    ("Telephone & Internet Charges", "Phone bills, internet, Wi-Fi"),
    ("Office Misc Expenses", "General office miscellaneous expenses"),
    ("Marketing & Ads", "Instagram, Google, social media ads, marketing expenses"),
    ("Registration & Subscription", "Software subscriptions, registrations, licenses"),
    ("Stationery", "Paper, pens, printer supplies"),
    ("Insurance Premium", "Insurance payments"),
    ("Administrative Overheads", "General admin expenses"),
    ("Local Conveyance", "Local travel, auto, cab for staff"),
    ("Business Developement", "Business development expenses"),
    ("Professional Charges", "CA, legal, consultant fees"),
    ("Donations", "Donations and charitable contributions"),
    ("Purchase @ 12%", "GST purchases at 12% rate"),
    ("Purchase @ 18%", "GST purchases at 18% rate"),
    ("Purchase @ 28%", "GST purchases at 28% rate"),
]

table_data = [["#", "Expense Ledger (Tally Head)", "When to Use"]]
for i, (name, desc) in enumerate(expense_ledgers, 1):
    table_data.append([str(i), name, desc])
add_table(table_data)

doc.add_page_break()


# ════════════════════════════════════════════════════════════════════
#  5. VENDOR / PARTY LEDGER LIST
# ════════════════════════════════════════════════════════════════════

doc.add_heading("5. Vendor / Party Ledger List", level=1)

doc.add_paragraph(
    "Use these exact names in the Vendor_Name column (Column F). "
    "These are the party (Sundry Creditor) ledgers in Tally."
)

vendors = [
    ("VENDOR for Non GST Purchases", "Default for all non-GST / petty cash Journal entries"),
    ("Cash", "Direct cash payments"),
    ("Axis Bank", "Bank payments via Axis Bank"),
    ("ICICI Bank", "Bank payments via ICICI Bank"),
    ("A2A Marketing & Service", "Marketing services vendor"),
    ("Akshay - Milk", "Milk supply"),
    ("Amazon Purchases", "Amazon online purchases"),
    ("BLINK COMMERCE PRIVATE LIMITED", "Blinkit orders"),
    ("Channakeshava Gas Cylinder Dealer", "Gas cylinder supply"),
    ("Classic Laundry", "Laundry services"),
    ("D Mart", "D Mart grocery purchases"),
    ("Flower Merchant", "Flower supply for decoration"),
    ("Ganga Bakers", "Bakery items"),
    ("GOOGLE INDIA PVT LTD", "Google ads / services"),
    ("Gopal Stores", "General store purchases"),
    ("Instagram", "Instagram advertising"),
    ("Meat Shop", "Meat / chicken supply"),
    ("Mynthra Shopping", "Myntra purchases"),
    ("Petrol Bunk", "Petrol / diesel"),
    ("PORTER", "Porter delivery / transport"),
    ("Prema Vehicle Hiring", "Vehicle rental"),
    ("Royal Mart", "Royal Mart purchases"),
    ("Seven Star Luxury Rentals", "Luxury rental services"),
    ("SHANTI FLOWER", "Flower supply"),
    ("SLK Cane Furniture", "Furniture purchases"),
    ("Softdrinks Vendor", "Soft drinks supply"),
    ("Sri Omkar Traders (MRA Vegetable)", "Vegetable supply"),
    ("Sri Venkateshwara Traders (Groceries)", "Grocery supply"),
    ("Swiggy Instamart", "Swiggy Instamart orders"),
    ("The Big Market", "Big Market grocery"),
    ("Y Muniraju Vegetable Agents", "Vegetable supply"),
]

table_data = [["#", "Vendor Name (exact)", "Type / Description"]]
for i, (name, desc) in enumerate(vendors, 1):
    table_data.append([str(i), name, desc])
add_table(table_data)

add_note("If a new vendor is not in this list, inform the accounts team to create the ledger in Tally first under Sundry Creditors group.", "warning")

doc.add_page_break()


# ════════════════════════════════════════════════════════════════════
#  6. DEPARTMENT LIST
# ════════════════════════════════════════════════════════════════════

doc.add_heading("6. Department List", level=1)

doc.add_paragraph("Use these in the Department column (Column C) for internal tracking:")

departments = [
    ("Thota Kitchen", "All kitchen and food-related expenses"),
    ("Thota Decor", "Decoration, flowers, event decoration"),
    ("Studio", "Studio operations"),
    ("Logistics", "Fuel, transport, vehicle expenses"),
    ("Staff Welfare", "Staff food, welfare, benefits"),
    ("Admin", "Administrative, office expenses"),
    ("Marketing", "Marketing, advertising, social media"),
    ("Maintenance", "Repairs, plumbing, electrical, AC"),
]

table_data = [["Department", "Covers"]]
for name, desc in departments:
    table_data.append([name, desc])
add_table(table_data)

doc.add_paragraph()
doc.add_heading("Paid By — Who Made the Payment:", level=2)

paid_by = ["Account", "Niveditha", "Akshata", "Raju", "Manu", "Krupa", "Harish", "Nisha", "Padum"]
doc.add_paragraph(", ".join(paid_by))

doc.add_page_break()


# ════════════════════════════════════════════════════════════════════
#  7. HOW TO UPLOAD — STEP BY STEP
# ════════════════════════════════════════════════════════════════════

doc.add_heading("7. How to Upload — Step by Step", level=1)

doc.add_heading("Step 1: Start the App", level=2)
steps = [
    "Make sure Tally Prime is running with THOTA HOSPITALITY LLP company open",
    "Go to the AGENT folder: D:\\RaviGowda\\Thota Documents\\AGENT\\",
    "Double-click START_APP.bat",
    "A black command window will open — DO NOT CLOSE IT",
    "Your browser will open automatically showing the upload page",
    "Check the top-right corner — it should show \"Tally Connected\" with a green dot",
]
for s in steps:
    doc.add_paragraph(s, style='List Bullet')

add_note("If the green dot is red (\"Tally Disconnected\"), check that Tally Prime is running and XML Server is enabled on port 9000.", "warning")

doc.add_paragraph()
doc.add_heading("Step 2: Upload Your Excel File", level=2)
steps = [
    "Click the upload area OR drag and drop your Excel file onto it",
    "The file must be .xlsx format",
    "Wait for validation to complete (usually 1-2 seconds)",
]
for s in steps:
    doc.add_paragraph(s, style='List Bullet')

doc.add_paragraph()
doc.add_heading("Step 3: Review the Validation", level=2)
steps = [
    "Check the summary boxes — Total Entries, Journal count, Purchase count, Total Amount",
    "GREEN message (\"All entries are valid\") → you can proceed",
    "RED messages (errors) → fix the errors in Excel and re-upload",
    "YELLOW messages (warnings) → review but you can still proceed",
    "Scroll through the data table to verify entries look correct",
]
for s in steps:
    doc.add_paragraph(s, style='List Bullet')

doc.add_paragraph()
doc.add_heading("Step 4: Enter Voucher Prefix", level=2)
doc.add_paragraph(
    "Type a prefix for the voucher numbers. This helps identify uploaded batches in Tally."
)
add_table([
    ["Month", "Suggested Prefix", "Voucher Numbers Created"],
    ["March 2026", "MAR26", "MAR26-001, MAR26-002, MAR26-003, ..."],
    ["April 2026", "APR26", "APR26-001, APR26-002, APR26-003, ..."],
    ["May 2026", "MAY26", "MAY26-001, MAY26-002, MAY26-003, ..."],
    ["Test upload", "TEST", "TEST-001, TEST-002, TEST-003, ..."],
])

add_note("If you leave the prefix empty, Tally will assign its own automatic numbers.", "tip")

doc.add_paragraph()
doc.add_heading("Step 5: Push to Tally", level=2)
steps = [
    "Click the green \"Push to Tally\" button",
    "Wait for the upload to complete (usually 5-30 seconds depending on number of entries)",
    "You will see a SUCCESS message with the count of vouchers created",
    "The last voucher number is shown (e.g., MAR26-150)",
]
for s in steps:
    doc.add_paragraph(s, style='List Bullet')

doc.add_paragraph()
doc.add_heading("Step 6: Verify in Tally", level=2)
doc.add_paragraph("See Section 8 below for how to verify the uploaded entries in Tally Prime.")

doc.add_page_break()


# ════════════════════════════════════════════════════════════════════
#  8. HOW TO VERIFY IN TALLY
# ════════════════════════════════════════════════════════════════════

doc.add_heading("8. How to Verify in Tally", level=1)

doc.add_heading("Check in Day Book:", level=2)
steps = [
    "Open Tally Prime → Gateway of Tally",
    "Go to Day Book (press D or click Day Book)",
    "Press F2 to change the date period",
    "Set From and To dates to match your uploaded data (e.g., 01-03-2026 to 31-03-2026)",
    "You should see all uploaded vouchers with your prefix (e.g., MAR26-001, MAR26-002...)",
    "Click on any voucher to open and verify details",
]
for i, s in enumerate(steps, 1):
    doc.add_paragraph(f"{i}. {s}", style='List Number')

doc.add_paragraph()
doc.add_heading("Check totals:", level=2)
steps = [
    "Go to Trial Balance (Gateway of Tally → Trial Balance)",
    "Look at expense ledgers (Thota Kitchen, Staff Welfare, etc.)",
    "The amounts should match what was in your Excel",
]
for s in steps:
    doc.add_paragraph(s, style='List Bullet')

doc.add_paragraph()
doc.add_heading("If something is wrong:", level=2)
steps = [
    "Open the incorrect voucher in Day Book",
    "Press Alt+D to delete the voucher",
    "Confirm deletion",
    "Fix the data in your Excel file",
    "Re-upload the corrected file",
]
for s in steps:
    doc.add_paragraph(s, style='List Bullet')

add_note("You can delete individual vouchers in Tally. There is no need to delete all and re-upload.", "tip")

doc.add_page_break()


# ════════════════════════════════════════════════════════════════════
#  9. COMMON ERRORS & FIXES
# ════════════════════════════════════════════════════════════════════

doc.add_heading("9. Common Errors & Fixes", level=1)

add_table([
    ["Error Message", "Cause", "Fix"],
    ["Date is missing", "Column A is empty for a row", "Enter the date in DD-MM-YYYY format"],
    ["Voucher_Type is missing", "Column B is empty", "Select Journal or Purchase from dropdown"],
    ["Invalid Voucher_Type", "Typed something other than Journal/Purchase", "Use exactly \"Journal\" or \"Purchase\""],
    ["Expense_Ledger is missing", "Column E is empty", "Select from dropdown — must match Tally exactly"],
    ["Vendor_Name is missing", "Column F is empty", "Select from dropdown — see Section 5 for list"],
    ["Invalid Amount", "Column G is empty, zero, or has text", "Enter a positive number (no commas, no ₹ symbol)"],
    ["Tally Disconnected", "Tally Prime is not running or XML Server is off", "Open Tally Prime, press F12 → Connectivity → port 9000"],
    ["Tally timeout", "Tally has a popup/dialog blocking it", "Switch to Tally, close any popup, press Esc to go to main screen"],
    ["Ledger not found", "Vendor or expense name doesn't exist in Tally", "Create the ledger in Tally first, or check spelling"],
])

doc.add_page_break()


# ════════════════════════════════════════════════════════════════════
#  10. FAQ
# ════════════════════════════════════════════════════════════════════

doc.add_heading("10. Frequently Asked Questions", level=1)

faqs = [
    (
        "How many entries can I upload at once?",
        "There is no practical limit. The tool has been tested with up to 1000+ entries in a single upload. "
        "However, uploading one month at a time (e.g., all of March) is recommended."
    ),
    (
        "Can I upload the same file twice?",
        "Yes, but it will create duplicate entries in Tally. Each upload creates new vouchers. "
        "If you accidentally upload twice, delete the duplicates from Tally's Day Book."
    ),
    (
        "What if I need to add a new vendor?",
        "First create the vendor ledger in Tally Prime (under Sundry Creditors group). "
        "Then you can type the exact vendor name in the Excel. Inform the admin to update the Excel template dropdown."
    ),
    (
        "What if I need to add a new expense head?",
        "First create the expense ledger in Tally Prime (under the appropriate group like Direct Expenses or Indirect Expenses). "
        "Then use the exact name in the Excel. Inform the admin to update the template."
    ),
    (
        "Can I upload from a different computer?",
        "Yes, as long as: (1) the AGENT folder is accessible (copy or share it), (2) Python is installed, "
        "(3) the computer can reach Tally on the network (update TALLY_HOST in .env file)."
    ),
    (
        "What happens if Tally is closed during upload?",
        "The upload will fail with a \"Connection Refused\" error. Simply start Tally, "
        "open the company, and try again."
    ),
    (
        "Do I need internet for this?",
        "No. Everything runs locally on your computer. No internet is required."
    ),
    (
        "The app URL (localhost:5050) is not opening?",
        "Make sure the black command window (START_APP.bat) is still running. "
        "If it closed, double-click START_APP.bat again."
    ),
    (
        "Can I use the same prefix for multiple uploads in a month?",
        "Yes, the numbering continues from where it left off within each upload. But if you do two separate uploads "
        "with the same prefix, numbering resets (e.g., MAR26-001 again). Use different prefixes like MAR26A, MAR26B "
        "to avoid confusion."
    ),
]

for q, a in faqs:
    p = doc.add_paragraph()
    run = p.add_run(f"Q: {q}")
    run.bold = True
    run.font.size = Pt(11)
    p2 = doc.add_paragraph(f"A: {a}")
    p2.paragraph_format.space_after = Pt(12)

doc.add_page_break()


# ════════════════════════════════════════════════════════════════════
#  QUICK REFERENCE CARD
# ════════════════════════════════════════════════════════════════════

doc.add_heading("Quick Reference Card", level=1)
doc.add_paragraph("Print this page and keep it at your desk.")

doc.add_paragraph()
p = doc.add_paragraph()
run = p.add_run("Daily Upload — 3 Easy Steps:")
run.bold = True
run.font.size = Pt(14)

doc.add_paragraph()

steps_summary = [
    "1️⃣  Double-click START_APP.bat  →  Browser opens",
    "2️⃣  Drag & drop your Excel file  →  Check green validation",
    "3️⃣  Enter prefix (e.g., MAR26)  →  Click Push to Tally  →  Done!",
]
for s in steps_summary:
    p = doc.add_paragraph(s)
    p.paragraph_format.space_after = Pt(8)
    for run in p.runs:
        run.font.size = Pt(13)

doc.add_paragraph()
p = doc.add_paragraph()
run = p.add_run("Remember:")
run.bold = True
p2 = doc.add_paragraph("• Tally Prime must be running before you start")
p3 = doc.add_paragraph("• Keep the black command window open while uploading")
p4 = doc.add_paragraph("• Check Day Book in Tally after upload to verify")
p5 = doc.add_paragraph("• If anything is wrong, delete individual vouchers with Alt+D in Tally")

doc.add_paragraph()
p = doc.add_paragraph()
run = p.add_run("App URL: ")
run.bold = True
run2 = p.add_run("http://localhost:5050")
run2.font.size = Pt(14)
run2.font.color.rgb = RGBColor(26, 35, 126)
run2.bold = True

p = doc.add_paragraph()
run = p.add_run("AGENT Folder: ")
run.bold = True
run2 = p.add_run("D:\\RaviGowda\\Thota Documents\\AGENT\\")
run2.font.name = 'Consolas'
run2.font.size = Pt(10)

p = doc.add_paragraph()
run = p.add_run("Support: ")
run.bold = True
run2 = p.add_run("Contact your system administrator or Ravi Gowda")


# ── Save ───────────────────────────────────────────────────────────

output_path = Path(r"D:\RaviGowda\Thota Documents\AGENT\.tmp\Thota_Tally_Upload_SOP.docx")
output_path.parent.mkdir(parents=True, exist_ok=True)
doc.save(str(output_path))
print(f"✅ SOP document created: {output_path}")
